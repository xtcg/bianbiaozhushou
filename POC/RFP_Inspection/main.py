from typing import Optional, List
from .retrieval import JinaEmbeddings
from langchain_core.documents import Document as langchain_Document
import json
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from .preprocessor import preprocess
from .preprocessor import timer
import re
# from concurrent.futures import ProcessPoolExecutor
from concurrent.futures import ThreadPoolExecutor
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Length, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING


embeddings = JinaEmbeddings(url='http://127.0.0.1:50086/api/v2/embed_onnx')
# base_url = 'https://vip-hk-s1.zeabur.app/v1'
api_key_official = "sk-proj-qiX4t372WKrmR5r0UxHbT3BlbkFJum2Dsi9T34ckZ0PGyz03"
base_url_official = None
api_key = 'sk-1UAoLHJ5EPoZn53v26C46770C6Cf44A6AbEa4d9523B6CaD2'
base_url = "https://vip-api-s1-global.aiearth.dev/v1"
# api_key = "sk-XsVovJaUjvMsduq6EbB30e8eD4714d678699A13f19B17350"
# api_key = 'sk-Skf3FAzpAeqWWgMA9aD1B9Ff33Ac4e318eC40eB3D9888327'
# base_url = 'https://api.wangyingzhe.cloud/v1'
client = OpenAI(api_key=api_key, base_url=base_url)
client_official = OpenAI(api_key=api_key_official, base_url=base_url_official)
k_section = 5
k = 5
section_num = 10

def print_to_file(message, file_name='./state_review.log'):
    with open(file_name, 'a') as file:
        print(message, file=file)

def extract_elements(type, response, results):
    content_match = re.search(r'"结论": "(.*?)"', response, re.DOTALL)
    reference_match = re.search(r'"依据": "(.*?)"', response, re.DOTALL)
    reason_match = re.search(r'"理由": "(.*?)"', response, re.DOTALL)
    #print("content", content_match)
    #print("reference", reference_match)
    results[type] = {}
    results[type]["content"] = content_match.group(1) if content_match else ""
    results[type]["reference"] = reference_match.group(1) if reference_match else ""
    results[type]["reason"] = reason_match.group(1) if reason_match else ""
    

@timer
def retrieve_content(type, db_zhaobiao_sections, db_toubiao_sections, text_splitter, query_dict, zhaobiao_query=None):
    print_to_file(f"开始检索“{type}”的内容...")
    if db_toubiao_sections:
        section_query = query_dict[type]['toubiao_section'] if query_dict.get(type, None) else type.replace("投标", "").replace("要求", "")
        if isinstance(section_query, list):
            toubiao_section_retriever = db_toubiao_sections.as_retriever(search_kwargs={'k': section_num//len(section_query)})
            toubiao_sections = []
            for q in section_query:
                sections_p = toubiao_section_retriever.invoke(q)
                toubiao_sections.extend(sections_p)
        else:
            toubiao_section_retriever = db_toubiao_sections.as_retriever(search_kwargs={'k': section_num})
            toubiao_sections = toubiao_section_retriever.invoke(section_query)
        toubiao_documents_by_section_split = [
            document
            for doc_section in toubiao_sections
            for document in text_splitter.create_documents([doc_section.metadata['data']], [{'section': doc_section.metadata['section']}])
        ]
        toubiao_db_by_section_split = FAISS.from_documents(documents=toubiao_documents_by_section_split, embedding=embeddings)
        toubiao_retriever = toubiao_db_by_section_split.as_retriever(search_kwargs={'k': query_dict[type].get("k", k) if query_dict.get(type, None) else k})
        toubiao_docs = toubiao_retriever.invoke(query_dict[type]['toubiao_content'] if query_dict.get(type, None) else type.replace("投标", "").replace("要求", ""))
        toubiao_docs.extend(toubiao_documents_by_section_split[:query_dict[type].get("k_section", k_section) if query_dict.get(type, None) else k_section])
        toubiao_result_dict = [{'section': doc.metadata['section'], 'data': doc.page_content} for doc in toubiao_docs]
        
        with open(f"rag_results_{type}_toubiao.json", "w") as f:
            json.dump(toubiao_result_dict, f, indent=4, ensure_ascii=False)
    else:
        toubiao_result_dict = None

    if db_zhaobiao_sections:
        section_query = query_dict[type]['zhaobiao_section'] if query_dict.get(type, None) else type
        if isinstance(section_query, list):
            zhaobiao_section_retriever = db_zhaobiao_sections.as_retriever(search_kwargs={'k': section_num//len(section_query)})
            zhaobiao_sections = []
            for q in section_query:
                sections_p = zhaobiao_section_retriever.invoke(q)
                zhaobiao_sections.extend(sections_p)
        else:
            zhaobiao_section_retriever = db_zhaobiao_sections.as_retriever(search_kwargs={'k': section_num})
            zhaobiao_sections = zhaobiao_section_retriever.invoke(section_query)
        zhaobiao_documents_by_section_split = [
            document
            for doc_section in zhaobiao_sections
            for document in text_splitter.create_documents([doc_section.metadata['data']], [{'section': doc_section.metadata['section']}])
        ]
        zhaobiao_db_by_section_split = FAISS.from_documents(documents=zhaobiao_documents_by_section_split, embedding=embeddings)
        zhaobiao_retriever = zhaobiao_db_by_section_split.as_retriever(search_kwargs={'k': query_dict[type].get("k", k) if query_dict.get(type, None) else k})
        zhaobiao_docs = zhaobiao_retriever.invoke(query_dict[type]['zhaobiao_content'] if query_dict.get(type, None) else type)
        zhaobiao_docs.extend(zhaobiao_documents_by_section_split[:query_dict[type].get("k_section", k_section) if query_dict.get(type, None) else k_section])
        zhaobiao_result_dict = [{'section': doc.metadata['section'], 'data': doc.page_content} for doc in zhaobiao_docs]
        with open(f"rag_results_{type}_zhaobiao.json", "w") as f:
            json.dump(zhaobiao_result_dict, f, indent=4, ensure_ascii=False)
    elif zhaobiao_query is not None:
        zhaobiao_result_dict = zhaobiao_query[type]
    else:
        zhaobiao_result_dict = None

    print_to_file(f"“{type}”的内容检索完毕！")
    return toubiao_result_dict, zhaobiao_result_dict


# types: 项目业主、项目名称、标段名称、工期、质量、安全、文中落款日期
def key_element_review(types, zhaobiao_path, toubiao_path, outline, complete_query_path: str = './RFP_Inspection/query.json', out_path="./result.docx"):
    print_to_file("开始进行关键内容审查...")
    results = {}
    zhaobiao_metadata_path = preprocess(docx_path=zhaobiao_path)
    toubiao_metadata_path = preprocess(docx_path=toubiao_path)
    with open(zhaobiao_metadata_path, 'r', encoding='utf-8') as file:
        zhaobiao_metadata = json.load(file)
    with open(toubiao_metadata_path, 'r', encoding='utf-8') as file:
        toubiao_metadata = json.load(file)
    print_to_file(f"开始创建检索...")
    # load complete query
    with open(complete_query_path, 'r', encoding='utf-8') as file:
        query_dict = json.load(file)
   
    # recursive character text splitter
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=100,
        length_function=len,
    )
    zhaobiao_sections = [langchain_Document(page_content="".join(section_data['section'][-2:]), metadata=section_data) for
                    section_data in zhaobiao_metadata]
    db_zhaobiao_sections = FAISS.from_documents(documents=zhaobiao_sections,
                                    embedding=embeddings)
    toubiao_sections = [langchain_Document(page_content="".join(section_data['section']), metadata=section_data) for
                    section_data in toubiao_metadata]
    db_toubiao_sections = FAISS.from_documents(documents=toubiao_sections,
                                    embedding=embeddings)
    
    # 使用with语句来确保线程池最后能够被正确关闭
    with ThreadPoolExecutor(max_workers=11) as executor:
        # 创建一个将type映射到相应future的字典
        future_to_type = {executor.submit(construct_section_content, type, db_zhaobiao_sections, db_toubiao_sections, text_splitter, query_dict): type for type in types}
        for future in future_to_type:
            type = future_to_type[future]
            response = future.result()
            extract_elements(type, response, results)
            
    with open("elements_key.json", "w") as f:
        json.dump(results, f, indent=4, ensure_ascii=False)     
    elements2docx(outline, element_doc_path="./elements_key.json", output_docx_path="./关键内容审查.docx")
    print_to_file("关键内容审查结束！")
    return zhaobiao_metadata, db_toubiao_sections, text_splitter, query_dict


def document_element_review(zhaobiao_metadata, db_toubiao_sections, text_splitter, query_dict):

    print_to_file("开始响应招标文件要求检查！")
    response = extract_evaluation_elements(zhaobiao_metadata)
    try:
        start = response.find("{")
        end = response.rfind("}") + 1
        json_str = response[start:end]
        evaluation_elements = json.loads(json_str)
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        evaluation_elements = {}
    outline_xiangying = {}
    for category, details in evaluation_elements.items():
        if isinstance(details, dict):
            outline_xiangying[category] = list(details.keys())
        else:
            continue
    
    rules_xiangying = {}
    for category, details in evaluation_elements.items():
        if isinstance(details, dict):
            rules_xiangying.update(details)
        else:
            continue
    
    types_xiangying = list(rules_xiangying.keys())

    results = {}
    if evaluation_elements:
        with ThreadPoolExecutor(max_workers=3) as executor:
            # 创建一个将type映射到相应future的字典
            future_to_type = {executor.submit(construct_section_content, type, None, db_toubiao_sections, text_splitter, query_dict, rules_xiangying): type for type in types_xiangying}
            for future in future_to_type:
                type = future_to_type[future]
                response = future.result()
                extract_elements(type, response, results)

    with open("elements_xiangying.json", "w") as f:
        json.dump(results, f, indent=4, ensure_ascii=False)
    elements2docx(outline_xiangying, element_doc_path="./elements_xiangying.json", output_docx_path="./响应招标文件要求审查.docx")
    print_to_file("响应招标文件要求检查结束！")     
    
    
@timer
def generate_response(type, toubiao_result_dict, zhaobiao_result_dict):
    print_to_file(f"开始生成“{type}”的内容...")
    sysprompt = f'''
你是招标公司审核员，你的工作是针对给定的<方面>，判断<投标文件>是否满足<招标文件>的要求，并给出对应的理由和索引。

请按照下面的格式进行输出，但不要抄袭所给例子中的内容：

{{
"结论": "符合",
"理由": "投标文件中关于项目名称的表述为'羊村青青草原一期200MW光伏电站项目'与'羊村青青草原一期200MW光伏电站项目 EPC 总承包'，关于在项目上的名称一致",
"依据": "招标文件部分 '一、投标函、投标函附录及投标人廉洁自律承诺书' 相关章节"
}}

{{
"结论": "符合",
"理由": "提瓦特公司三藩市布里瓦尔大道 50 万千瓦项目'与'提瓦特公司三藩市布里瓦尔大道 50 万千瓦项目 EPC 总承包'项目名称一致",   
"依据": "第六章 施工组织设计、（一）工程概况、3项目简介 第7页"
}}

{{
"结论": "不符合", 
"理由": "项目名称长沙200兆瓦/300兆瓦时储能电站项目错误，应为滨海200兆瓦/400兆瓦时储能电站项目。 " ,
"依据": "第六章 施工组织设计、（一）工程概况、3项目简介 第7页 "
}}

现在需要你来完成：

<方面>
{type}
</方面>
'''
    userprompt = f'''
<投标文件>
{toubiao_result_dict}
</投标文件>

<招标文件>
{zhaobiao_result_dict}
</招标文件>

请注意！"理由"必须严格按照下面的固定格式进行输出：
"招标文件中的{type}为xxx；投标文件中的{type}为xxx；分析：xxx，符合/不符合招标文件中关于{type}的要求。"

"依据"是指相关内容在招投标文件所出现的页码或章节。

现在请你开始输出：
'''
    message = [
        {"role": "system", "content": sysprompt},
        {"role": "user", "content": userprompt}
    ]
    model = 'gpt-4-turbo-preview'

    try:
        stream = client.chat.completions.create(
            model=model, messages=message,
            stream=True,
            max_tokens=4095,
            temperature=1,
            top_p=1
        )

        response = ""
        for chunk in stream:
            if not chunk.choices:
                continue
            chunk_response = chunk.choices[0].delta.content
            if chunk_response:
                response += chunk_response
            # yield response
        print_to_file(f"“{type}”的内容生成完毕！")
        return response
    
    except Exception as e:
        print('generation error while', type)
        print(f"An unexpected error occurred: {e}")
        print_to_file(f"“{type}”生成过程错误！")
        return ""

@timer
def construct_section_content(type, db_zhaobiao_sections, db_toubiao_sections, text_splitter, query_dict, zhaobiao_query=None):
    toubiao_result_dict, zhaobiao_result_dict = retrieve_content(type, db_zhaobiao_sections, db_toubiao_sections, text_splitter, query_dict, zhaobiao_query)
    response = generate_response(type, toubiao_result_dict, zhaobiao_result_dict)
    print(response)
    return response

@timer
def extract_evaluation_elements(zhaobiao_metadata):
    print_to_file(f"提取评审信息...")
    data = [None, None]
    for dp in zhaobiao_metadata:
        if (dp['section'][-1] == "评标办法前附表") or ("3.1 初步评审" in dp['section'][-1]):
            data[0] = dp['data']
        if dp['section'][-1] == "投标人须知前附表":
            data[1] = dp['data']
    if data:
        sysprompt = f'''
你是一个投标文件撰写者，你需要对招标文件的<评审标准>进行提取，以便于从投标文件中检索到相关信息。

<评审标准>
{data[0]}
</评审标准>
    '''
        userprompt = f'''
<投标人须知>
{data[1]}
</投标人须知>

------------------------------------------------------------------

你需要严格将输出内容整理成以下json的形式，务必不要增加额外的层级关系：
```json
{{ 评审因素  : {{评审标准或对应条款: 内容完整的具体标准}}, ...}},
```
评审因素可以是“符合性审查标准”、“形式评审标准”、“响应性评审标准”等。具体标题需要通过<评审标准>中提到的内容来确定。如果<评审标准>没有提到任何相关内容，则直接采用“评审标准”作为评审因素。

请注意将评审标准中关于<投标人须知>的要求与所提供材料的内容相对应。输出中不要出现“"符合第二章“投标人须知”第 3.2 款规定” 之类的语句。

请注意<评审标准>中与评审内容无关的条款，例如候选人数量、评分、计算公式等不需要进行输出。
审查点无需提取评分部分，无需出现商务、技术、报价的详细评审。

请务必输出包含结尾的完整内容：
    '''

        message = [
        {"role": "system", "content": sysprompt},
        {"role": "user", "content": userprompt}
    ]
        model = 'gpt-4-1106-preview'

        # print('---------------------------\n')
        # print(sysprompt)
        # print('\n--------------------------\n')
        # print(userprompt)
        # print('\n---------------------------')
        # exit()

        try:
            stream = client_official.chat.completions.create(
                model=model, messages=message,
                stream=True,
                max_tokens=4095,
                temperature=0.85,
                top_p=1
            )

            response = ""
            for chunk in stream:
                if not chunk.choices:
                    continue
                chunk_response = chunk.choices[0].delta.content
                if chunk_response:
                    response += chunk_response
                # yield response
            print(response)
            return response
        
        except Exception as e:
            print(f"An unexpected error occurred: {e}")
            print('generation error')
            print_to_file(f"生成过程错误！")
            return None
            
    else:
        print_to_file("未找到评标办法前附表")
        return None



def elements2docx(outline, element_doc_path="./elements.json", output_docx_path="./审查.docx"):
    with open(element_doc_path, 'r', encoding='utf-8') as file:
        element_dict = json.load(file)

    doc = Document()

    style_title = doc.styles.add_style('style_title', WD_STYLE_TYPE.PARAGRAPH)
    style_title.font.name = 'Arial'
    style_title.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    style_title.font.size = Pt(15)
    style_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_title.paragraph_format.line_spacing = 1.5

    style_level1 = doc.styles.add_style('style_level1', WD_STYLE_TYPE.PARAGRAPH)
    style_level1.base_style = doc.styles['Heading 1']  # 基于现有的标题1样式
    style_level1.font.name = 'Arial'  # 拉丁文本字体
    style_level1.element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')  # 亚洲文本字体
    style_level1.font.bold = True
    style_level1.font.size = Pt(14)
    style_level1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style_level1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_level1.paragraph_format.outline_level = 0

    style_level2 = doc.styles.add_style('style_level2', WD_STYLE_TYPE.PARAGRAPH)
    style_level2.font.name = 'Arial'
    style_level2.font.size = Pt(14)
    style_level2.font.bold = True
    style_level2.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
    style_level2.paragraph_format.line_spacing = 1.5  # 1.5行间距
    style_level2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_level2.paragraph_format.outline_level = 1  # 大纲级别为2时，outline_level设为1

    style_element = doc.styles.add_style('style_element', WD_STYLE_TYPE.PARAGRAPH)
    style_element.base_style = doc.styles['Normal']
    style_element.font.name = 'Arial'
    style_element.element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
    style_element.font.bold = True
    style_element.font.size = Pt(14)
    style_element.paragraph_format.line_spacing = 1.5
    style_element.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_element.paragraph_format.first_line_indent = Inches(0.5)
    
    style = doc.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.first_line_indent = Inches(0.5)

    style_reference = doc.styles.add_style('style_reference', WD_STYLE_TYPE.PARAGRAPH)
    style_reference.font.name = 'Arial'
    style_reference.element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
    style_reference.font.underline = True
    style_reference.font.size = Pt(14)
    style_reference.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style_reference.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_reference.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_paragraph("审查结果", style='style_title')
    for section, subsections in outline.items():
        doc.add_paragraph(section, style='style_level1')
        if subsections:
            if isinstance(subsections, list):
                for subsection in subsections:
                        doc.add_paragraph(subsection, style='style_level2')
                        if element_dict.get(subsection, False):
                            doc.add_paragraph(element_dict[subsection]["content"], style='style')
                            doc.add_paragraph(element_dict[subsection]["reason"], style='style')
                            doc.add_paragraph(element_dict[subsection]["reference"], style='style_reference')
    
    doc.save(output_docx_path)

if __name__ == '__main__':
    
    types = [
    "项目名称",
    # "项目业主", "标段名称",
    "工期",
      "质量", "安全", 
    # "文中落款日期"
    ]
    
    outline = {
    "关键内容审查": 
        ["项目名称","工期", "质量", "安全"]
        }

    zhaobiao_metadata, db_toubiao_sections, text_splitter, query_dict = key_element_review(types, 'RFP_Inspection/提瓦特-招标文件脱敏稿.docx', 'RFP_Inspection/提瓦特-技术标脱敏稿.docx', outline=outline)
    document_element_review(zhaobiao_metadata, db_toubiao_sections, text_splitter, query_dict)
    # run_parallel_tasks(types, './test2.docx', outline, out_path="./abstract2.docx")
    # run_parallel_tasks(types, './test3.docx', outline, out_path="./abstract3.docx")

    # type = "支付方式"
    # with open(f'./rag_results_{type}.json', 'r', encoding='utf-8') as file:
    #     result_dict = json.load(file)
    # with open('./prompt.json', 'r', encoding='utf-8') as file:
    #     prompt_dict = json.load(file)
    
    # print(prompt_dict[type]["sysprompt"])
    # print('\n\n---------------------------\n\n'),
    # print(prompt_dict[type]["userprompt"].format(rag_content=result_dict))

    