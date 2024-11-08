from typing import Optional, List
from .retrieval import JinaEmbeddings
from langchain_core.documents import Document as langchain_Document
import json
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from .preprocessor import preprocess
from .preprocessor import timer
import re
# from concurrent.futures import ProcessPoolExecutor
from concurrent.futures import ThreadPoolExecutor
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Length, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
import concurrent.futures


embeddings = JinaEmbeddings(url='http://127.0.0.1:50086/api/v2/embed_onnx')
# base_url = 'https://vip-hk-s1.zeabur.app/v1'
# api_key = 'sk-1UAoLHJ5EPoZn53v26C46770C6Cf44A6AbEa4d9523B6CaD2'
# base_url = "https://vip-api-s1-global.aiearth.dev/v1"
# api_key = "sk-proj-qiX4t372WKrmR5r0UxHbT3BlbkFJum2Dsi9T34ckZ0PGyz03"
api_key = "sk-proj-0xgTHzIK7bgJFVGCiisZT3BlbkFJ66f7e7IqQx9oKsK6pIuV"
base_url = None
# api_key = "sk-kM2oyhZKyEfwf6EMAb0966De3b7649A481343c0fD8592076"
# api_key = "sk-XsVovJaUjvMsduq6EbB30e8eD4714d678699A13f19B17350"
# api_key = 'sk-Skf3FAzpAeqWWgMA9aD1B9Ff33Ac4e318eC40eB3D9888327'
# base_url = 'https://api.wangyingzhe.cloud/v1'
client = OpenAI(api_key=api_key, base_url=base_url)
k_section = 6
k = 10
section_number = 10
sysprompt_template = "你现在是一个招标书的审阅者。请你对标书内容进行摘要，提取标书中的“{generation_type}”的相关信息。\n\n输出的形式如下：\n<output>\n{{\n\"content\": ,\n\"reference\": \n}}\n</output>\n\n\n以下是给你的参考内容列表："
userprompt_template = "{rag_content}\n\n请注意，你需要按照下面的例子来理解例子中对输出内容的格式的要求，但不要抄袭例子里的内容：\n<output>\n{{\n\"content\": \"（1）承包人应按照结算款项金额向发包人提供符合税务规定的增值税专用发票，发包人在收到承包人提供的合格增值税专用发票后支付款项。\n（2）承包人应确保增值税专用发票真实、规范、合法，如承包人虚开或提供不合格的增值税专用发票，造成发包人经济损失的，承包人承担全部赔偿责任，并重新向发包人开具符合规定的增值税专用发票。\n（3）合同变更如涉及增值税专用发票记载项目发生变化的，应当约定作废、重开、补开、红字开具增值税专用发票。如果收票方取得增值税专用发票尚未认证抵扣，收票方应及时退回原发票，则可以由开票方作废原发票，重新开具增值税专用发票；如果原增值税专用发票已经认证抵扣，则由开票方就合同增加的金额补开增值税专用发票，就减少的金额依据收票方提供的红字发票信息表开具红字增值税专用发票。\",\n\"reference\": \"第一章 招标公告 第 9 页 \"\n}}\n</output>\n\n请充分利用参考内容列表中的信息，以章节（1）（2）的形式对内容进行组织梳理，输出跟例子格式相近的信息，要求每个章节有一定的篇幅，优先输出专用合同条款的内容。请注意，请提取全部参考内容中section为{generation_type}的data，然后将这些内容全部完整输出，并保证输出结尾内容。如果材料中没有相关信息，则输出空字符串“”："


def find_substring(string, substring_list):
    for substring in substring_list:
        if substring in string:
            return substring
    return None  # 如果没有找到，返回None

def print_to_file(message, file_name='./state_abstract.log'):
    with open(file_name, 'a') as file:
        print(message, file=file)

def extract_elements(type, response, results):
    content_match = re.search(r'"content": "(.*?")', response, re.DOTALL)
    reference_match = re.search(r'"reference": "(.*?)"', response, re.DOTALL)
    #print("content", content_match)
    #print("reference", reference_match)
    results[type] = {}
    content = content_match.group(1).replace('\\n', '\n') if content_match else ""
    if content.endswith('"'):
        results[type]["content"] = content[:-1]
    else:
        results[type]["content"] = content
    results[type]["reference"] = reference_match.group(1) if reference_match else ""
    

@timer
def retrieve_content(type: str, db_by_section, query_dict, text_splitter, sections=None):
    if sections is None:
        print_to_file(f"开始检索“{type}”的内容...")
        section_retriever = db_by_section.as_retriever(search_kwargs={'k': query_dict[type].get("section_number", section_number)})
        section_query = query_dict[type].get("section_query" ,query_dict[type]['query'])
        # print('section_query', section_query)
        if isinstance(section_query, list):
            sections = []
            for q in section_query:
                sections_p = section_retriever.invoke(q)
                sections.extend(sections_p)
        else:
            sections = section_retriever.invoke(section_query)
        documents_by_section_split = [
            document
            for doc_section in sections
            for document in text_splitter.create_documents([doc_section.metadata['data']], [{'section': doc_section.metadata['section'], 'page': doc_section.metadata['page']}])
        ]

        # need to be tested if needed
        for doc in documents_by_section_split:
            doc.page_content = f'{doc.metadata["section"][-1]}\n{doc.page_content}'

        try:
            db_by_section_split = FAISS.from_documents(documents=documents_by_section_split, embedding=embeddings)
            global_retriever = db_by_section_split.as_retriever(search_kwargs={'k': query_dict[type].get("k", k) if query_dict.get(type, False) else k})
            docs = global_retriever.invoke(query_dict[type]['query'] if query_dict.get(type, False) else type)
            docs.extend(documents_by_section_split[:query_dict[type].get("k_section", k_section) if query_dict.get(type, False) else k_section])
            result_dict = [{'section': doc.metadata['section'], 'page': doc.metadata['page'], 'data': doc.page_content} for doc in docs]
            with open(f"rag_results_{type}.json", "w") as f:
                json.dump(result_dict, f, indent=4, ensure_ascii=False)
            print_to_file(f"“{type}”的内容检索完毕！")
            return result_dict
        except:
            print_to_file(f"“{type}”的内容检索出错！")
            return None

    else:
        documents_by_section_split = [
            document
            for doc_section in sections
            for document in text_splitter.create_documents([doc_section.metadata['data']], [{'section': doc_section.metadata['section'], 'page': doc_section.metadata['page']}])
        ]

        # need to be tested if needed
        for doc in documents_by_section_split:
            doc.page_content = f'{doc.metadata["section"][-1]}\n{doc.page_content}'

        try:
            db_by_section_split = FAISS.from_documents(documents=documents_by_section_split, embedding=embeddings)
            global_retriever = db_by_section_split.as_retriever(search_kwargs={'k': query_dict[type].get("k", k) if query_dict.get(type, False) else k})
            docs = global_retriever.invoke(query_dict[type]['query'] if query_dict.get(type, False) else type)
            docs.extend(documents_by_section_split[:query_dict[type].get("k_section", k_section) if query_dict.get(type, False) else k_section])
            result_dict = [{'section': doc.metadata['section'], 'page': doc.metadata['page'], 'data': doc.page_content} for doc in docs]
            with open(f"rag_results_{type}.json", "w") as f:
                json.dump(result_dict, f, indent=4, ensure_ascii=False)
            print_to_file(f"“{type}”的内容检索完毕！")
            return result_dict
        except:
            print_to_file(f"“{type}”的内容检索出错！")
            return None


@timer
def extract_section_name(metadata):
    section_string = ['合同']
                    #   ,  '条款']

    print_to_file(f"开始提取合同价格与支付相关章节...")
    relevant_sections = []
    for dp in metadata:
        section_query = " ".join(dp['section'])
        if all(field in section_query for field in section_string):
            relevant_sections.append(dp)
 
    docs = [langchain_Document(page_content=" ".join(dp['section']), metadata={'data': dp['data'], 'page': dp['page'], 'section': dp['section']}) for dp in relevant_sections]
    db = FAISS.from_documents(documents=docs, embedding=embeddings)
    retriever = db.as_retriever(search_kwargs={'k': 40})
    sections = retriever.invoke('专用 合同价格与支付')
    section_list = [section.page_content for section in sections]

    sysprompt = """
现在你需要对招标文件进行摘要任务。

你需要首先确定合同价格与支付的子章节，子章节的确认需要以专用合同为主。然后你需要提供该章节对应的原本招标文件的章节。并以下面的dict形式输出，但不要抄袭例子中的内容：

```json
{{
"element": "竣工结算",
"source": "第四章 合同条款及格式 第二节专用合同条款 17．合同价格与支付 17.5竣工结算"
}}
```
"""
    userprompt = f"{section_list}\n\n注意\"source\"必须完整，与给的章节名完全一致！\n请从下面章节中抽取出与合同价格与支付相关的子章节，并进行输出，注意子章节之间不能重复，不要输出与合同价格与支付无关的信息。如果没有，则输出空列表："

    print('==============================================')
    print(userprompt)
    print('==============================================')
    print(sysprompt)
    print('==============================================')


    message = [
        {"role": "system", "content": sysprompt},
        {"role": "user", "content": userprompt}
    ]
    model = 'gpt-4-turbo-preview'

    try:
        stream = client.chat.completions.create(
            model=model, messages=message,
            stream=True,
            max_tokens=4095,
            temperature=1,
            top_p=1
        )

        response = ""
        for chunk in stream:
            if not chunk.choices:
                continue
            chunk_response = chunk.choices[0].delta.content
            if chunk_response:
                response += chunk_response
            # yield response
        print('==============================================')
        print(response)
        print('==============================================')

    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        print('generation error')
        print_to_file(f"提取过程错误！")
        return None, None, None

    try:
        start = response.find("[")
        end = response.rfind("]") + 1
        json_str = response[start:end]
        generation_elements = json.loads(json_str)
        print('generation_elements', generation_elements)
    except:
        start = response.find("{")
        end = response.rfind("}") + 1
        json_str = response[start:end]
        generation_elements = json.loads(json_str)
        print('generation_elements', generation_elements)

    section_dict = {}
    for g in generation_elements:
        # element_list.append(g['element'])
        section_dict[g['source']] = g['element']
    
    element_list = []
    sections_selected= {}
    section_list = list(section_dict.keys())
    for d in docs:
        relevant_element = find_substring(d.page_content, section_list)
        if relevant_element:
            # print('1', d.page_content)
            element_list.append(section_dict[relevant_element])
            if section_dict[relevant_element] in list(section_dict.keys()):
                sections_selected[section_dict[relevant_element]].extend([langchain_Document(page_content=d.metadata['data'], metadata={'section': d.page_content, 'data': d.metadata['data'], 'page': d.metadata['page']})])
            else:
                sections_selected[section_dict[relevant_element]] = [langchain_Document(page_content=d.metadata['data'], metadata={'section': d.page_content, 'data': d.metadata['data'], 'page': d.metadata['page']})]

    element_list = list(set(element_list))
    return element_list, sections_selected, docs


def run_parallel_tasks(types, doc_path, outline, complete_query_path: str = './RFP_assistant/query.json', pdf_path=None, out_path="./abstract.docx"):
    print_to_file("任务开始！")
    results = {}
    meta_data_path = preprocess(docx_path=doc_path, pdf_path=pdf_path)
    print_to_file(f"开始创建检索...")
    # load complete query
    with open(complete_query_path, 'r', encoding='utf-8') as file:
        query_dict = json.load(file)
    # load metadata
    with open(meta_data_path, 'r', encoding='utf-8') as file:
        metadata = json.load(file)[1:]
    # recursive character text splitter
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=100,
        length_function=len,
    )
    documents_sections = [langchain_Document(page_content=section_data['section'][-1], metadata=section_data) for
                    section_data in metadata]
    db_by_section = FAISS.from_documents(documents=documents_sections,
                                    embedding=embeddings)

    # 使用with语句来确保线程池最后能够被正确关闭
    with ThreadPoolExecutor(max_workers=11) as executor:
        # 创建一个将type映射到相应future的字典
        future_to_type = {executor.submit(construct_section_content, type, db_by_section, query_dict, text_splitter): type for type in types}
        for future in future_to_type:
            type = future_to_type[future]
            response = future.result()
            if response:
                extract_elements(type, response, results)

    element_list, sections_selected, sections_relevant = extract_section_name(metadata)

    print('==============================================')
    print(element_list)
    print('==============================================')
    print(sections_selected)
    print('==============================================')
    if element_list and sections_selected:
        with ThreadPoolExecutor(max_workers=11) as executor:
            # 创建一个将type映射到相应future的字典
            future_to_type = {executor.submit(construct_section_content, type, db_by_section, query_dict, text_splitter, sections_selected[type]): type for type in element_list}
            for future in future_to_type:
                type = future_to_type[future]
                response = future.result()
                if response:
                    extract_elements(type, response, results)
        outline["七、合同价格与支付"] = element_list
    
    elif sections_relevant:
        response = construct_section_content('合同价格与支付', db_by_section, query_dict, text_splitter, sections_relevant)
        if response:
            extract_elements('合同价格与支付', response, results)
            outline["七、合同价格与支付"] = ['合同价格与支付']

    with open("elements.json", "w") as f:
        json.dump(results, f, indent=4, ensure_ascii=False)     
    elements2docx(outline, output_docx_path=out_path)

    print_to_file("任务结束！")
    
@timer
def generate_response(type, result_dict, prompt_path='./RFP_assistant/prompt.json'):
    print_to_file(f"开始生成“{type}”的内容...")
    with open(prompt_path, 'r', encoding='utf-8') as file:
        prompt_dict = json.load(file)
    
    sysprompt = prompt_dict[type]["sysprompt"] if prompt_dict.get(type, False) else sysprompt_template.format(generation_type=type)
    userprompt = prompt_dict[type]["userprompt"].format(rag_content=result_dict) if prompt_dict.get(type, False) else userprompt_template.format(rag_content=result_dict, generation_type=type)
    message = [
        {"role": "system", "content": sysprompt},
        {"role": "user", "content": userprompt}
    ]
    model = 'gpt-4-turbo-preview'

    # print('--------------------------\n')
    # print(sysprompt)
    # print('\n-------------------------\n')
    # print(userprompt)
    # print('\n--------------------------')
    # exit()

    try:
        stream = client.chat.completions.create(
            model=model, messages=message,
            stream=True,
            max_tokens=4095,
            temperature=0.8,
            top_p=1
        )

        response = ""
        for chunk in stream:
            if not chunk.choices:
                continue
            chunk_response = chunk.choices[0].delta.content
            if chunk_response:
                response += chunk_response
            # yield response
        print_to_file(f"“{type}”的内容生成完毕！")
        return response
    
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        print('generation error while', type)
        print_to_file(f"“{type}”生成过程错误！")
        return ""


def construct_section_content(type, db_by_section, query_dict, text_splitter, sections=None):
    rag_results = retrieve_content(type, db_by_section, query_dict, text_splitter, sections=sections)
    if rag_results:
        response = generate_response(type, rag_results)
        print(response)
        return response
    else:
        return None

    

def elements2docx(outline, element_doc_path="./elements.json", output_docx_path="./abstract.docx"):
    with open(element_doc_path, 'r', encoding='utf-8') as file:
        element_dict = json.load(file)

    doc = Document()

    style_title = doc.styles.add_style('style_title', WD_STYLE_TYPE.PARAGRAPH)
    style_title.font.name = 'Arial'
    style_title.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    style_title.font.size = Pt(15)
    style_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_title.paragraph_format.line_spacing = 1.5

    style_level1 = doc.styles.add_style('style_level1', WD_STYLE_TYPE.PARAGRAPH)
    style_level1.base_style = doc.styles['Heading 1']  # 基于现有的标题1样式
    style_level1.font.name = 'Arial'  # 拉丁文本字体
    style_level1.element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')  # 亚洲文本字体
    style_level1.font.bold = True
    style_level1.font.size = Pt(14)
    style_level1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style_level1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_level1.paragraph_format.outline_level = 0

    style_level2 = doc.styles.add_style('style_level2', WD_STYLE_TYPE.PARAGRAPH)
    style_level2.font.name = 'Arial'
    style_level2.font.size = Pt(14)
    style_level2.font.bold = True
    style_level2.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
    style_level2.paragraph_format.line_spacing = 1.5  # 1.5行间距
    style_level2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_level2.paragraph_format.outline_level = 1  # 大纲级别为2时，outline_level设为1

    style_element = doc.styles.add_style('style_element', WD_STYLE_TYPE.PARAGRAPH)
    style_element.base_style = doc.styles['Normal']
    style_element.font.name = 'Arial'
    style_element.element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
    style_element.font.bold = True
    style_element.font.size = Pt(14)
    style_element.paragraph_format.line_spacing = 1.5
    style_element.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_element.paragraph_format.first_line_indent = Inches(0.5)
    
    style = doc.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.first_line_indent = Inches(0.5)

    style_reference = doc.styles.add_style('style_reference', WD_STYLE_TYPE.PARAGRAPH)
    style_reference.font.name = 'Arial'
    style_reference.element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
    style_reference.font.underline = True
    style_reference.font.size = Pt(14)
    style_reference.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style_reference.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_reference.paragraph_format.first_line_indent = Inches(0.5)

    if element_dict.get('项目名称', False):
        doc.add_paragraph(f"{element_dict['项目名称']['content']}\n商务分析", style='style_title')
    for section, subsections in outline.items():
        doc.add_paragraph(section, style='style_level1')
        if subsections:
            if isinstance(subsections, list):
                for subsection in subsections:
                        doc.add_paragraph(subsection, style='style_level2')
                        if element_dict.get(subsection, False):
                            for line in element_dict[subsection]["content"].split('\n'):
                                doc.add_paragraph(line, style='style')
                            doc.add_paragraph(element_dict[subsection]["reference"], style='style_reference')
            elif isinstance(subsections, dict):
                for subsection, elements  in subsections.items():
                    doc.add_paragraph(subsection, style='style_level2')
                    for element in elements:
                        doc.add_paragraph(element, style='style_element')
                        if element_dict.get(element, False):
                            for line in element_dict[element]["content"].split('\n'):
                                doc.add_paragraph(line, style='style')
                            doc.add_paragraph(element_dict[element]["reference"], style='style_reference')
    
    doc.save(output_docx_path)

if __name__ == '__main__':
    
    types = [
    # "项目名称",
    #   "项目业主",  
    # "项目概况",
    # "建设地点", 
    # "计划工期", 
    # "质量标准",
    #  "最高投标限价",
    #   "投标文件递交",
    #     "招标范围",
    # "合同价格",
    #  "预付款", "工程进度付款",
    #   "质量保证金", "竣工结算",
    # "最终结清", "增值税专用发票",
    #  "支付方式"
]
    
    outline = {
    "一、项目概况与招标范围": {
        "1. 项目概况": ["项目名称", "项目业主", "项目背景","项目概况","其他需重点关注的技术要求",
            "建设地点", "质量标准","最高投标限价","投标文件递交"],
        "2. 招标范围" : ["招标范围"]
    },
    # "二、资信及业绩要求": [],
    # "三、主要人员资格要求、财务要求、信誉要求": [],
    # "四、踏勘现场": [],
    # "五、最高投标限价": [],
    # "六、评标方法": [],
    "七、合同价格与支付": [],
    # "八、罚则": [],
    # "九、其他合同条件": [],
    # "十、投标报价": []
    }

    run_parallel_tasks(types, doc_path='./RFP_assistant/HZZFCG-2023-135杭州市数据资源管理局任务在线项目_原件202306182229.docx', pdf_path='./RFP_assistant/HZZFCG-2023-135杭州市数据资源管理局任务在线项目_原件202306182229.pdf' ,outline=outline, out_path="./abstract_test.docx")
    # run_parallel_tasks(types, './RFP_assistant/test2.docx', outline, out_path="./abstract2.docx")
    # run_parallel_tasks(types, './RFP_assistant/test3.docx', outline, out_path="./abstract3.docx")

    # type = "预付款"
    # with open(f'./rag_results_{type}.json', 'r', encoding='utf-8') as file:
    #     result_dict = json.load(file)
    # with open('./prompt.json', 'r', encoding='utf-8') as file:
    #     prompt_dict = json.load(file)
    
    # print(prompt_dict[type]["sysprompt"])
    # print('\n\n---------------------------\n\n'),
    # print(prompt_dict[type]["userprompt"].format(rag_content=result_dict))


    # section_string = ['合同',  '条款']

    # with open(f'./RFP_assistant/test3.json', 'r', encoding='utf-8') as file:
    #     result_dict = json.load(file)

    # results = []
    # for dp in result_dict:
    #     section_query = " ".join(dp['section'])
    #     if all(field in section_query for field in section_string):
    #         results.append(section_query)
    
    # print(results)
    # docs = [langchain_Document(page_content=i) for i in results]
    # db = FAISS.from_documents(documents=docs,
    #                                 embedding=embeddings)
    # retriever = db.as_retriever(search_kwargs={'k': 40})
    # o = retriever.invoke('专用 合同价格与支付')
    # print(o)
    






    