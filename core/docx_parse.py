import bisect
from docx import Document
import re
import json
from typing import List
import os 
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from pathlib import Path
import sys

POTENTIAL_PATTERNS = [
                     r'(\（[一二三四五六七八九十]{1,2}\）)',
                     r'(\（[1234567890]{1,2}\）)',
                     r'(\([一二三四五六七八九十]{1,2}\))',
                     r'(\([1234567890]{1,2}\))',
                     r'(\d+\.\d+\.\d+\.)[^\.\d]+',
                     r'(\d+\.\d+\.\d+)[^\.\d]+',
                     r'(\d+\.\d+\.)[^\.\d]+',
                     r'(\d+\.\d+)[^\.\d]+',
                     r'(\d+\.)[^\.\d]+',
                     r'([一二三四五六七八九十]{1,2}\.).*',
                     r'([一二三四五六七八九十]{1,2}\、).*',
                     r'(第[1234567890]{1,2}[节部分])',
                     r'(第[一二三四五六七八九十]{1,2}[节部分])[^，。；：,.;:]*',
                     r'(第[1234567890]{1,2}章)[^，。；：,.;:]*',
                     r'(第[一二三四五六七八九十]{1,2}章)[^，。；：,.;:]*',
                     r'(第[1234567890]{1,2}卷).*',
                     r'(第[一二三四五六七八九十]{1,2}卷).*',
                     r'(第[IVX]{1,3}卷).*',
                     r'(第[\u2160-\u2169]卷).*'
                    ]

class Docx():
    def __init__(self, file_path: str, use_cache: bool = False):
        """
        contents: read all the available contents in the file, mainly processed in function docx2json
        cached_dir: the direction to save some results
        rag: the final splitted and chunked paragraphs, saved as 'rag.json' in cached_dir
        use_cache: use stored results to speed up the process
        catalogs: save the contents of catalogs from the file
        """
        assert file_path.endswith('docx'), "input should be a docx file"
        self.doc = Document(file_path)
        self.file_path = file_path
        self.contents = None
        # self.cached_dir = os.path.join('/home/ron/jiamin/bianbiaozhushou/data/',self.file_path.split('/', -1)[-1].replace('.docx', "").strip())
        self.cached_dir = Path(__file__).parent.parent.joinpath('data/test_outline').joinpath(self.file_path.split('/', -1)[-1].replace('.docx', "").strip())
        self.hash_dir = None
        if not os.path.exists(self.cached_dir):
            print(self.cached_dir)
            os.makedirs(self.cached_dir)
        self.rag = list()
        self.rag1 = list()
        self.use_cache = use_cache
        self.catalogs = list()
        self.head = ''
        self.head_of_content = 0

    def docx2json(self,use_cache=False):
        """
        read from file to self.contents
        """
        des_file = os.path.join(self.cached_dir,'content.json')
        if self.hash_dir and os.path.exists(os.path.join(self.hash_dir,'content.json')):
            self.contents = None
            with open(os.path.exists(os.path.join(self.hash_dir,'content.json')), 'r') as f:
                self.contents = json.load(f)
                if 'catalogs' in self.contents:
                    self.catalogs = self.contents['catalogs']
                    self.rag.append(self.contents['append_to_rag'])
            return 
        self.contents = {
                         "paragraphs": {},
                         "tables": {},
                         "check":{},
                         "catalog_para":{
                                        "exist": False,
                                        "index": None,
                                        },
                         "catalog_table":{
                                        "exist": False,
                                        "index": None,
                                        },
                        }
        para_point = 0
        table_point = 0
        key_index = 0
        #count the number of paragraphs and tables
        pattern = re.compile(r'目.*录\t*')
        for index, element in enumerate(self.doc.element.body):
            if element.tag.endswith('p'):
                #current element is a paragraph
                origin = self.doc.paragraphs[para_point].text
                splitted = origin.split('\n')
                if len(splitted)>1:
                    splitted = [i+'\t' for i in splitted]
                for paragraph in splitted:
                    if len(paragraph)>1:
                        if (not self.contents['catalog_para']['exist']) and pattern.match(paragraph):
                            self.contents['catalog_para']['exist'] = True                                                             
                            self.contents['catalog_para']['index'] = key_index
                        self.contents['paragraphs'][str(key_index)] = {'text': paragraph,
                                                            'index': str(index),}
                        self.contents['check'][str(key_index)] = {
                                                            "type": 'paragraph',
                                                            "index": para_point
                                                            }
                        key_index += 1
                para_point += 1
            elif element.tag.endswith('tbl'):
                #current element is a table
                table = self.doc.tables[table_point]
                cur = ''
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]    
                    if "最高投标限价" in ''.join(row_data) and row_data[-1] == '/':
                        row_data[-1] = row_data[-1][:-1]+'未设置最高投标报价'
                    cur += ' '.join(row_data) + '\n'
                if (not self.contents['catalog_table']['exist']) and pattern.match(table.rows[0].cells[0].text):
                    self.contents['catalog_table']['exist'] = True
                    self.contents['catalog_table']['index'] = key_index
                self.contents['tables'][str(key_index)] = {"text": cur,
                                                      "index":str(index),}
                self.contents['check'][str(key_index)] = {
                                                    "type": 'table',
                                                    "index": table_point
                                                    }
                table_point += 1
                key_index += 1
            elif element.tag.endswith('sdt'):
                #current element is a catalog table
                def get_sdt_content(element, namespaces):
                    """
                    提取sdt元素中的文本内容
                    """
                    sdt_content = element.find('.//w:sdtContent', namespaces)
                    if sdt_content is not None:
                        paragraphs = sdt_content.findall('.//w:p', namespaces)
                        text_content = []
                        for paragraph in paragraphs:
                            texts = [node for node in paragraph.itertext()]
                            index = [len(x) for x in texts]
                            text_content.append(texts[index.index(max(index))].split('\t')[0].strip())
                    return   text_content
                namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                content = get_sdt_content(element,namespaces)
                # print(content)
                para_point -= 1
                for c in content:
                    self.contents['paragraphs'][str(key_index)] = {"text": c+'\t', "index":str(index)}
                    self.contents['check'][str(key_index)] = {
                                                            "type": 'paragraph',
                                                            "index": para_point
                                                            }
                    if (not self.contents['catalog_para']['exist']) and pattern.match(c):
                        self.contents['catalog_para']['exist'] = True                                                             
                        self.contents['catalog_para']['index'] = key_index
                    key_index += 1
                para_point += 1
                # if self.contents['catalog_para']['exist'] == False and self.contents['catalog_table']['exist'] == False:
                #     self.catalogs[:] = content[:]
                #     self.rag.append({
                #         'text': content,
                #         'index': [index],
                #         'first_index': str(index),
                #         'final_index': str(index),
                #         'length': 0,
                #         'type': '目录',
                #     })
                #     self.contents['catalogs'] = self.catalogs
                #     self.contents['append_to_rag'] = self.rag[0]
        with open(des_file,'w') as f:
            f.write(json.dumps(self.contents,indent=2,ensure_ascii=False)+'\n')
        if self.hash_dir and not os.path.exists(os.path.join(self.hash_dir,'content.json')):
            with open(os.path.join(self.hash_dir, 'content.json'),'w') as f:
                f.write(json.dumps(self.contents,indent=2,ensure_ascii=False)+'\n')
        return


    def get_rag(self):
        self.rag = []
        self.rag1 = []
        self.catalogs = []
        des_file = os.path.join(self.cached_dir,'rag.json')
        des_file1 = os.path.join(self.cached_dir,'rag1.json')
        self.docx2json(self.use_cache)
        self.find_cover()
        self.find_catalogs()
        self.find_contents()
        if os.path.exists(des_file) and os.path.exists(des_file1) and self.use_cache:
            with open(des_file) as f:
                self.rag = json.load(f)
            with open(des_file1) as f:
                self.rag1 = json.load(f)
            return self.rag
        self.post_process()
        self.post_process1()
        self.insert_table(self.rag)
        self.insert_table(self.rag1)
        with open(des_file, 'w') as f:
            f.write(json.dumps(self.rag,indent=2,ensure_ascii=False)+'\n')
        with open(des_file1, 'w') as f:
            f.write(json.dumps(self.rag1,indent=2,ensure_ascii=False)+'\n')
        return self.rag


    def find_cover(self):
        """
        find the cover of the file, the key is location of the keyword "目录"
        during docx2json, record the position for catalogs, and use bisect to locate the index in the index list
        index list is not continous
        """
        index_list = list(map(int, self.contents['paragraphs'].keys()))
        temp =  {
            'text': [],
            'index': [],
            'first_index': 0,
            'final_index': 0,
            'length': 0,
            'type': '封面',
        }
        if len(self.catalogs):
            # print(1)
            # print(len(index_list),self.rag)
            #when there is a catalog table, and self.catalogs is handled in function docx2json
            pos = bisect.bisect_left(index_list, int(self.rag[0]['first_index']))
            temp['text'][:] = [remove_space(self.contents['paragraphs'][str(x)]['text']) for x in index_list[:pos]]
            temp['index'][:] = [x for x in index_list[:pos]]
        elif self.contents['catalog_para']['exist'] and self.contents['catalog_para']['index'] < 100:
            #when catalogs in paragraph object
            pos = bisect.bisect_left(index_list, self.contents['catalog_para']['index'])
            temp['text'][:] = [remove_space(self.contents['paragraphs'][str(x)]['text']) for x in index_list[:pos]]
            temp['index'][:] = [x for x in index_list[:pos]]
        elif self.contents['catalog_table']['exist'] and self.contents['catalog_table']['index'] < 100:
            #when catalogs in table object
            pos = bisect.bisect_left(self.contents['catalog_table']['index'])
            temp['text'][:] = [remove_space(self.contents['paragraphs'][str(x)]['text']) for x in index_list[:pos]]
            temp['index'].append(self.contents['catalog_table']['index'])
        else:
            temp =  {
            'text': [],
            'index': [],
            'first_index': 0,
            'final_index': 0,
            'length': 0,
            'type': 'content',
        }
            temp['text'] = [self.contents['paragraphs'][x]['text'] for x in self.contents['paragraphs']]
            temp['index'] = [x for x in self.contents['paragraphs']]
            self.catalogs = temp['text']
        if len(temp['index']):
            temp['first_index'] = temp['index'][0]
            temp['final_index'] = temp['index'][-1]
            temp['length'] = sum([len(x) for x in temp['text']])
            self.rag.append(temp)
        self.rag.reverse()
        name_from_path = remove_space(self.file_path.split('/')[-1].replace('.docx',''))
        count = 10000
        for t in self.contents['paragraphs']:
            cur = self.minDistance(name_from_path, self.contents['paragraphs'][t]['text'])
            if cur < count:
                count = cur
                self.head = self.contents['paragraphs'][t]['text']
        return 

    def find_catalogs(self):
        """
        find catalogs contents, match the specific pattern like r".*\t\d*"
        """
        if len(self.catalogs):
            #already read from catalog table
            return
        content = list()
        pattern = re.compile(r".*\t\d*")
        pattern1 = re.compile(r".*∙+\d*")
        index_list = list(map(int, self.contents['paragraphs'].keys()))
        temp =  {
            'text': [],
            'index': [],
            'first_index': 0,
            'final_index': 0,
            'length': 0,
            'type': '目录',
        }
        pos = -1
        mode = 0
        if self.contents['catalog_para']['exist']:
            pos = bisect.bisect_left(index_list, self.contents['catalog_para']['index'])
            content = self.contents['paragraphs']
        elif self.contents['catalog_table']['exist']:
            pos = bisect.bisect_left(index_list, self.contents['catalog_table']['index'])
            content = self.contents['tables'][[pos]]
            mode = 1
        else:
            return
        
        if mode:
            cur = content['text'].split('\n')
            temp['text'][:] = cur[:]
            self.catalogs = cur[:]
            temp['index'].append(pos)
        else:
            temp['text'].append(remove_space(content[str(index_list[pos])]['text'].split('\t')[0]))
            temp['index'].append(index_list[pos])
            self.catalogs.append(remove_space(temp['text'][-1]))
            pattern121 = None
            flag = True
            count = 0
            while flag and count < 5:
                pos += 1
                for s in POTENTIAL_PATTERNS[::-1]:
                    pattern121 = re.compile(s)
                    if r:= pattern121.match(content[str(index_list[pos])]['text']):
                        flag = False
                        break
                if flag:
                    temp['text'].append(clean_text(content[str(index_list[pos])]['text'].split('\t')[0]))
                    temp['index'].append(index_list[pos])
                    count += 1
            while pos < len(index_list):               
                text = content[str(index_list[pos])]['text']
                splitted = text.split('\t')
                if pattern121.match(text) and clean_text(text) in self.catalogs:
                    if len(temp['index']):
                        temp['first_index'] = temp['index'][0]
                        temp['final_index'] = temp['index'][-1]
                        temp['length'] = sum([len(x) for x in temp['text']])
                        self.rag.append(temp)
                    return
                if not splitted[-1].isdigit():
                    text = ''.join(splitted)
                else:
                    text = ''.join(splitted[:-1])
                splitted = text.split('∙')
                if not splitted[-1].isdigit():
                    text = ''.join(splitted)
                else:
                    text = ''.join(splitted[:-1])
                #remove '\t' and special character in the text
                # if pattern.match(content[str(index_list[pos])]['text']) or pattern1.match(content[str(index_list[pos])]['text']):
                temp['text'].append(remove_space(text))
                temp['index'].append(index_list[pos])
                # print(temp['text'][-1])
                self.catalogs.append(remove_space(temp['text'][-1]))
                pos += 1    
        return 

    def find_contents(self):
        """
        read contents from the file, split the contents according to the catalogs
        """
        index_list = list(map(int, self.contents['paragraphs'].keys()))
        temp =  {
            'text': [],
            'index': [],
            'first_index': 0,
            'final_index': 0,
            'length': 0,
            'type': 'content',
        }
        #final position of catalogs
        last_index = self.rag[-1]['final_index'] if len(self.rag) else 0
        point = bisect.bisect_left(index_list, int(last_index)) + 1
        if point == len(index_list):
            point = 0
        self.rag1.append({'text': [clean_text(self.contents['paragraphs'][str(x)]['text']) for x in index_list[point:]],
                          'index': index_list[point:],
                          'first_index': index_list[point:][0],
                          'final_index': index_list[point:][-1],
                          'length': sum([len(x) for x in [self.contents['paragraphs'][str(x)]['text'] for x in index_list[point:]]]),
                          'type': 'content',})
        seen = set(map(remove_space,self.catalogs))
        self.catalogs = list(map(remove_space,self.catalogs))
        point = bisect.bisect_left(index_list, int(last_index)) + 1
        if len(self.catalogs) == 0:
            temp['text'].extend([remove_space(self.contents['paragraphs'][str(x)]['text']) for x in index_list[point:]])
            temp['index'].extend(index_list[point:])
            temp['first_index'] = temp['index'][0]
            temp['final_index'] = temp['index'][-1]
            temp['length'] = sum([len(x) for x in temp['text']])
            self.rag.append(temp)
            self.head_of_content = temp['first_index']
            return 

        pre = ''
        while point < len(index_list):
            cur_pos = index_list[point]
            text = clean_text(self.contents['paragraphs'][str(cur_pos)]['text'])
            if text in seen:
                if len(temp['text']) and not text == pre[:len(text)]:
                    temp['final_index'] = temp['index'][-1]
                    temp['length'] = sum([len(x) for x in temp['text']])
                    temp['first_index'] = temp['index'][0]
                    temp['type'] = 'content'
                    self.rag.append(temp)
                    if self.head_of_content == 0:
                        self.head_of_content = temp['first_index']
                pre = text
                temp =  {
                            'text': [],
                            'index': [],
                            'first_index': 0,
                            'final_index': 0,
                            'length': 0,
                            'type': 'content',
                        }
            temp['text'].append(text)
            temp['index'].append(cur_pos)
            point += 1
        
        if len(temp['text']):
            temp['final_index'] = temp['index'][-1]
            temp['length'] = sum([len(x) for x in temp['text']])
            temp['first_index'] = temp['index'][0]
            temp['type'] = 'content'
            self.rag.append(temp)
        return 

    def post_process(self):
        """
        try to find more shorter chunks
        """
        ans = []
        ans[:] = self.rag[:]
        count = 0
        point = len(POTENTIAL_PATTERNS) - 4
        while point > 0:
            res = []
            for con in ans:
                if con['length'] <= 2048 or con['text'][0].startswith('附件') or con['text'][0].startswith('附表') or con['text'][0].startswith('附录') or con['text'][0].startswith('目录'):
                    res.append(con)
                else:
                    temp = self.find_more_chunk(con,point)
                    res.extend(temp)
            point -= 1
            ans[:] = res[:]
        count = 0 
        while True:
            count += 1
            if count == 2:
                break
            res = []
            for con in ans:             
                if con['type'] == '目录' or con['type'] == '封面':
                    res.append(con)
                    continue
                if int(con['length']) <= 2048 or con['text'][0].startswith('附件') or con['text'][0].startswith('附表') or con['text'][0].startswith('附录'):
                    res.append(con)
                else:
                    temp = self.find_more_chunk1(con)
                    res.extend(temp)
            ans[:] = res[:]
        res = []
        for con in ans:
            if con['length'] > 512:
                temp = self.find_more_chunk2(con)
                res.extend(temp)
            else:
                res.append(con)
        self.rag[:] = res[:]
        self.rag[:] = ans[:]
        return
    
    def post_process1(self):
        """
        try to find more shorter chunks
        """
        ans = []
        ans[:] = self.rag1[:]
        res = []
        point = len(POTENTIAL_PATTERNS) - 1
        while point >= len(POTENTIAL_PATTERNS) - 4:
            res = []
            for con in ans:             
                if con['type'] == '目录' or con['type'] == '封面':
                    res.append(con)     
                temp = self.find_more_chunk(con, point)
                res.extend(temp)
            # print(len(ans),len(res))
            if len(ans) != len(res):
                break
            else:
                point -= 1
        ans = []
        for r in res:
            if r['length']>=50:
                ans.append(r)
        res = []
        point = len(POTENTIAL_PATTERNS) - 4
        while point >= 0:
            res = []
            for con in ans:             
                if con['type'] == '目录' or con['type'] == '封面':
                    res.append(con)     
                temp = self.find_more_chunk(con, point)
                res.extend(temp)
            # print(len(ans),len(res))
            if len(ans) != len(res):
                break
            else:
                point -= 1
        if len(ans) == len(res):        
            res = []
            for con in ans:             
                if con['type'] == '目录' or con['type'] == '封面':
                    res.append(con)
                    continue       
                temp = self.find_more_chunk1(con, used_in_rag=False)
                res.extend(temp)
        self.rag1 = []
        for r in range(len(res)):
            if res[r]['length'] < 800:
                continue
            else:
                self.rag1.extend(res[r:])
                break
        return

    def insert_table(self, rag_to_insert):
        """
        insert table contents to the chunks
        """
        check_list = [int(x['final_index']) for x in rag_to_insert]
        look_up = dict()
        for cont in rag_to_insert:
            look_up[int(cont['final_index'])] = cont
        for t in self.contents['tables']:
            insert_con = bisect.bisect_left(check_list, int(t))
            if insert_con == len(check_list):
                insert_con -= 1
            if insert_con < 0:
                    continue
            insert_index = list(map(int, look_up[check_list[insert_con]]['index']))
            insert_pos =  bisect.bisect_left(insert_index, int(t))
            if  insert_pos == 0:
                #table will not be the beginning
                insert_con -= 1
                if insert_con < 0:
                    continue
                insert_index = list(map(int, look_up[check_list[insert_con]]['index']))
                insert_pos =  bisect.bisect_left(insert_index, int(t))
            look_up[check_list[insert_con]]['text'] = look_up[check_list[insert_con]]['text'][:insert_pos]\
                                                    +[self.contents['tables'][t]['text']]\
                                                    +look_up[check_list[insert_con]]['text'][insert_pos:]
            look_up[check_list[insert_con]]['index'] = look_up[check_list[insert_con]]['index'][:insert_pos]\
                                                    +[int(t)]\
                                                    + look_up[check_list[insert_con]]['index'][insert_pos:]
        rag_to_insert = list()
        for i in look_up:
            rag_to_insert.append(look_up[i])
        return 

    def set_cached_dir(self, path: str, hash_path: str):
        self.cached_dir = path
        self.hash_dir = hash_path
        if not os.path.exists(self.cached_dir):
            os.makedirs(self.cached_dir)
        if not os.path.exists(self.hash_dir):
            os.makedirs(self.hash_dir)

    def _strQ2B(self, ustring):
        """把字符串全角转半角"""
        rstring = ""
        for uchar in ustring:
            inside_code = ord(uchar)
            if inside_code == 0x3000:
                inside_code = 0x0020
            else:
                inside_code -= 0xfee0
            if inside_code < 0x0020 or inside_code > 0x7e:  # 转完之后不是半角字符返回原来的字符
                rstring += uchar
            else:
                rstring += chr(inside_code)
        return rstring

    def minDistance(self, s: str, t: str) -> int:
        n, m = len(s), len(t)
        f = [[0] * (m + 1) for _ in range(n + 1)]
        f[0] = list(range(m + 1))
        for i, x in enumerate(s):
            f[i + 1][0] = i + 1
            for j, y in enumerate(t):
                f[i + 1][j + 1] = f[i][j] if x == y else \
                                    min(f[i][j + 1], f[i + 1][j], f[i][j]) + 1
        return f[n][m]

    def find_more_chunk(self, con: dict, pattern_point) -> list:
        """
        chunk from the start and find peer chunk, chunk based on the highest chapter level in a long paragraph
        """
        res = []
        pre = con['text'][0]
        temp =   {
            'text': [],
            'index': [],
            'first_index': 0,
            'final_index': 0,
            'length': 0,
            'type': '',
        }
        pattern = re.compile(POTENTIAL_PATTERNS[pattern_point])
        if not pattern.match(con['text'][0]):
            return [con]
        # print(POTENTIAL_PATTERNS[pattern_point])
        # print(POTENTIAL_PATTERNS[pattern_point])
        point = 0
        pre = con['text'][0]
        temp['text'].append(con['text'][point])
        temp['index'].append(con['index'][point])
        temp['type'] = con['type']
        point = 1
        pre = con['text'][0]
        while point < len(con['text']):
            text = clean_text(con['text'][point])
            if r:= pattern.match(text) and text != pre and text == remove_space(''.join(con['text'][point].split('\t'))):
                # print(text,pre, remove_space(con['text'][point]))
                temp['final_index'] = temp['index'][-1]
                temp['first_index'] = temp['index'][0]
                temp['length'] = sum([len(x) for x in temp['text']])
                temp['type'] = con['type']
                res.append(temp)
                temp =  {
                        'text': [],
                        'index': [],
                        'first_index': 0,
                        'final_index': 0,
                        'length': 0,
                        'type': '',
                    }
                pre = text
            temp['text'].append(text)
            temp['index'].append(con['index'][point])
            point += 1
        temp['final_index'] = temp['index'][-1]
        temp['first_index'] = temp['index'][0]
        temp['length'] = sum([len(x) for x in temp['text']])
        temp['type'] = con['type']
        res.append(temp)
        # print(len(res))
        return res

    def find_more_chunk1(self, con: dict, used_in_rag: bool = True) -> list:
        """
        chunk from the start + 1 and find children chunk, chunk based on the highest chapter level in a long paragraph
        """
        res = []
        pre = con['text'][0]
        cur = ''
        temp = {
                'text': [],
                'index': [],
                'first_index': 0,
                'final_index': 0,
                'length': 0,
                'type': '',
            }
        point = 0
        pre = con['text'][0]
        temp['text'].append(con['text'][point])
        temp['index'].append(con['index'][point])
        point += 1
        tag = True
        while point < len(con['text']) and tag:
            # if con['text'][point] == '1.招标条件':
                # print(121212,POTENTIAL_PATTERNS[:-4][::-1])
            for s in POTENTIAL_PATTERNS[:-4][::-1]:
                pattern = re.compile(s)
                if r:= pattern.match(con['text'][point]):
                    cur = s
                    # print(cur)
                    tag = False
                    break
            if tag == False:
                break
            temp['text'].append(con['text'][point])
            temp['index'].append(con['index'][point])
            point += 1
        # print(cur)
        temp['final_index'] = temp['index'][-1]
        temp['first_index'] = temp['index'][0]
        temp['length'] = sum([len(x) for x in temp['text']])
        temp['type'] = con['type']
        if used_in_rag:
            res.append(temp)
        if cur == '':           
            return [con]
        else:
            pattern = re.compile(cur)
        temp = {
                'text': [],
                'index': [],
                'first_index': 0,
                'final_index': 0,
                'length': 0,
                'type': '',
            }
        
        pre = con['text'][point]
        # print(pre)
        while point < len(con['text']):
            text = con['text'][point]
            if r:= pattern.match(text) and pre != text:
                temp['final_index'] = temp['index'][-1]
                temp['first_index'] = temp['index'][0]
                temp['length'] = sum([len(x) for x in temp['text']])
                temp['type'] = con['type']
                res.append(temp)
                temp = {
                        'text': [],
                        'index': [],
                        'first_index': 0,
                        'final_index': 0,
                        'length': 0,
                        'type': '',
                    }
                pre = text
                # print(pre)
            temp['text'].append(text)
            temp['index'].append(con['index'][point])
            point += 1
        temp['final_index'] = temp['index'][-1]
        temp['first_index'] = temp['index'][0]
        temp['length'] = sum([len(x) for x in temp['text']])
        temp['type'] = con['type']
        res.append(temp)
        return res

    def find_more_chunk2(self, con: dict) -> list:
        res = []
        n = len(con['text'])
        point = 0
        temp = {
                'text': [],
                'index': [],
                'first_index': 0,
                'final_index': 0,
                'length': 0,
                'type': '',
            }
        while point < n:
            count = 0
            pre = point
            while point < n and (len(temp['text']) == 0 or count + len(con['text'][point]) < 512):
                temp['text'].append(con['text'][point])
                temp['index'].append(con['index'][point])
                count += len(con['text'][point])
                point += 1
            temp['final_index'] = temp['index'][-1]
            temp['first_index'] = temp['index'][0]
            temp['length'] = sum([len(x) for x in temp['text']])
            temp['type'] = con['type']
            res.append(temp)
            temp = {
                'text': [],
                'index': [],
                'first_index': 0,
                'final_index': 0,
                'length': 0,
                'type': '',
            }
            count = 0
            while point < n and point > 0 and (count == 0 or count + len(con['text'][point]) < 100):
                count += len(con['text'][point])
                point -= 1
            if point == pre:
                point += 1
        return res



def remove_space(s: str):
    s = ''.join([x for x in s.split(' ')])
    # s = ''.join([x for x in s.split('\t')])
    return s

def clean_text(text: str) -> str:
    # 只保留英文字母、数字、中文字符和英文标点符号
    cleaned_text = re.sub(r"[^\w\s\u4e00-\u9fff.,!?;:\"'()\-]‘“’。，》《？》（）【】「」、", "", text)
    # 去掉多余的空格字符
    cleaned_text = re.sub(r"\s*", "", cleaned_text)
    cleaned_text = re.sub(r"；。", "", cleaned_text)
    return cleaned_text.strip()

def wirte2json(file_path: str, contents: List[dict]):
    with open(file_path, 'w', encoding='utf-8') as f:
        temp = []
        for c in contents:
            temp.append({
                "page": c['type'] if c['type'] == '目录' or c['type'] == '封面' else str(len(temp)-1),
                "content": '\n'.join(c['text']),
            })
        f.write(json.dumps(temp,indent=2, ensure_ascii = False)+'\n')
    return 

if __name__ == '__main__':
    # file_path = '/home/ron/jiamin/bianbiaozhushou/data/test_docx/招标文件-大唐博罗公庄镇 150MW 复合型光伏发电项目.docx'
    file_path = '/home/ron/jiamin/bianbiaozhushou/data/test_docx/内蒙古库布齐沙漠鄂尔多斯中北部新能源基地700万千瓦光伏项目工程设计施工采购.docx'
    file_path = "/home/ron/jiamin/bianbiaozhushou/data/test_docx/(招标文件)贵州华电望谟蔗香望南一期150MW农业光伏电站项目EPC总承包.docx"
    # file_path = '/home/ron/jiamin/bianbiaozhushou/data/temp/666729c5fd4046e79e17bbebb459f6e9/censor_key/bid.docx'
    file_path = '/home/ron/jiamin/bianbiaozhushou/data/test_docx/招标文件—浙江建德抽水蓄能电站项目设计采购施工EPC招标文件 20240724（终）.docx'
    file_path = '/home/ron/jiamin/bianbiaozhushou/data/test_docx/（招标文件）龙源电力招远市龙源新能源有限公司阜山镇对脚岭光伏项目EPC总承包.docx'
    print(clean_text("第一章招标公告；"))
    # doc = Document(file_path)
    # print(len(doc.element.body))
    # print(len(doc.part.rels.values()))
    # doc = Document(file_path)
    # para_point = 0
    # from lxml import etree
    # def get_paragraph_xml(paragraph):
    # # 获取段落的xml内容
    #     xml_str = paragraph._element.xml
    # # 使用lxml解析xml字符串
    #     xml_element = etree.fromstring(xml_str)
    # # 格式化输出
    #     return etree.tostring(xml_element, pretty_print=True, encoding='unicode')
    
    # # print(get_paragraph_xml(doc.paragraphs[155]))
    # for index, element in enumerate(doc.element.body):
    #         if element.tag.endswith('p'):
    #             #current element is a paragraph
    #             paragraph = doc.paragraphs[para_point].text
    #             para_point += 1
    #             if paragraph == '投标人资格要求':
    #                 print(get_paragraph_xml(doc.paragraphs[154]))
    #                 print(paragraph)
    #                 print(para_point)
    # doc = Docx(file_path, use_cache=False)
    # # doc.set_cached_dir('/home/ron/jiamin/bianbiaozhushou/data/temp/666729c5fd4046e79e17bbebb459f6e9/censor_key')

    # # for p in doc.doc.paragraphs:
    # #     if p.text.endswith('招标文件的获取'):
    # #         print(p.text)
    # #         for run in p.runs:
    # #             print(run.text)
    # # for index, element in enumerate(doc.doc.element.body):
    # #     if element.tag.endswith('textbox'):
    # #         print(element)
    # # doc.docx2json()
    # doc.get_rag()
    # # print(doc.head)
    # # a = sorted([x['length'] for x in doc.rag])
    # # print(a[:-5])
    # a = ['a','b','c']
    # print(f'{a}')
    # pattern = re.compile(r'(\d*\.)[^\d]+')
    # if pattern.match('1.招标条件'):
    #     print(11)
    