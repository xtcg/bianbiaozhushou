#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# @Time    : 2024/7/12 14:03
# @File    : parser.py
# @Software: PyCharm
import logging
import weakref
import typing as t

from docx import Document
from docx.parts.image import ImagePart
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.drawing import Drawing

logger = logging.getLogger(__name__)
P = t.TypeVar("P", bound=t.Union[Paragraph, Table, ImagePart])
# Paragraph style优先级
STYLE_PRIORITY = {
    "Title": 0,
    **{"Heading {i}": i for i in range(1, 10)},
    "Caption": 10,
    "Normal": 11,
    "List Paragraph": 12,
    "Body Text": 13,
    "Plain Text": 14,

    "Table Normal": 16,
    **{"toc {i}": 20 + i for i in range(1, 9)},
}


def get_paragraph_style_priority(p: Paragraph) -> int:
    if p.style.name in STYLE_PRIORITY:
        return STYLE_PRIORITY[p.style.name]


def table_to_markdown(table):
    markdown_output = ""
    markdown_table = []
    for row in table.rows:
        markdown_rows = []
        for cell in row.cells:
            # Strip any leading/trailing whitespace
            cell_text = cell.text.strip().replace("\n", "<br>")
            markdown_rows.append(cell_text)
        if any(markdown_rows):
            markdown_table.append("|" + "|".join(markdown_rows) + "|")

    # Append the table to the output, add a header separator
    header_separator = "|" + "|".join("---" for _ in table.columns) + "|"
    markdown_output += "\n".join(markdown_table[:1])  # Table headers
    markdown_output += "\n" + header_separator + "\n"
    markdown_output += "\n".join(markdown_table[1:])  # Table body
    # markdown_output += "\n\n"  # Add space between tables
    return markdown_output


def find_paragraph_extra(p: Paragraph) -> t.List:
    """
    Find images in paragraph
    """
    extras = []
    for run in p.runs:
        for content in run.iter_inner_content():
            if isinstance(content, Drawing):
                if image_part := draw_image_from_drawing(content):
                    extras.append(image_part)
    return extras


def draw_image_from_drawing(draw: Drawing) -> t.Optional[ImagePart]:
    blip = draw._element.xpath('.//a:blip')  # noqa
    if not blip:
        return
    embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')  # noqa
    if not embed:
        return
    image_part: t.Optional[ImagePart] = draw.part.related_parts.get(embed, None)
    return image_part


def parse_doc(doc: Document):
    last_p = None
    for s_index, section in enumerate(doc.sections):
        for p_index, p in enumerate(section.iter_inner_content()):
            if isinstance(p, Paragraph):
                if not p.text.strip():
                    continue
                yield ParsedDocument(p, s_index, p_index)
            elif isinstance(p, Table):
                yield ParsedDocument(p, s_index, p_index)
            elif isinstance(p, ImagePart):
                yield ParsedDocument(p, s_index, p_index)


class DocumentManager(dict):
    root: "ParsedDocument"

    def __init__(self, doc: Document):
        super().__init__()
        # self.doc = doc
        self._parse_doc(doc)

    def extract_toc(self):
        """
        extract toc from doc
        normally, the toc is the second section of the doc
        """

    def _parse_doc(self, doc):
        # 使用栈来实现
        # 根据标题级别来判断是否是子节点
        # 如果是子节点，就将其加入到父节点的children中并入栈
        # 如果不是子节点，出栈
        stack = []
        for s_index, section in enumerate(doc.sections):
            for p_index, p in enumerate(section.iter_inner_content()):
                if isinstance(p, Paragraph):
                    if not p.text.strip():
                        continue
                    parsed_doc = ParsedDocument(p, s_index, p_index)
                    if not stack:
                        self.root = parsed_doc
                    else:
                        stack[-1].add_child(parsed_doc)
                    stack.append(parsed_doc)
                elif isinstance(p, Table):
                    parsed_doc = ParsedDocument(p, s_index, p_index)
                    if not stack:
                        self.root = parsed_doc
                    else:
                        stack[-1].add_child(parsed_doc)
                    stack.append(parsed_doc)
                elif isinstance(p, ImagePart):
                    parsed_doc = ParsedDocument(p, s_index, p_index)
                    if not stack:
                        self.root = parsed_doc
                    else:
                        stack[-1].add_child(parsed_doc)
                    stack.append(parsed_doc)

    def get_paragraph(self, section: int, index: int) -> t.Optional["ParsedDocument"]:
        return self.get((section, index), None)


class ParsedDocument:
    __slot__ = ["p", "section", "index", "origin"]

    origin: t.List[P]

    header: str

    def __init__(self, p: P, section: int, index: int):
        self.p = weakref.ref(p)
        self.set_origins(p)
        self.section = section
        self.index = index

    def set_origins(self, p):
        self.origin.extend(find_paragraph_extra(p))

    @property
    def parent(self) -> t.Optional["ParsedDocument"]:
        if hasattr(self, "_parent"):
            return self._parent()
        return None

    @parent.setter
    def parent(self, value: "ParsedDocument"):
        self._parent = weakref.ref(value)  # noqa
        value.add_child(self)

    @property
    def children(self) -> t.List["ParsedDocument"]:
        if hasattr(self, "_children"):
            return self._children
        return []

    def add_child(self, child: "ParsedDocument"):
        if not hasattr(self, "_children"):
            self._children = []  # noqa
        self._children.append(child)


class ParsedParagraph(ParsedDocument):

    @property
    def text(self) -> str:
        if not hasattr(self, "_text"):
            self._text = self.p().text  # noqa
        return self._text


class ParsedTable(ParsedDocument):

    @property
    def text(self) -> str:
        if not hasattr(self, "_text"):
            self._text = table_to_markdown(self.p())  # noqa
        return self._text
