import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Length, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
import re


def json2docx(outline: dict, title: str, element_json: dict, docx_path: str, doc_type: str):
    """
    outline 是输出的docx，格式为
        "一、项目概况与招标范围":  ["项目名称", "项目业主", "项目背景","项目概况","其他需重点关注的技术要求",
                "建设地点", "质量标准","最高投标限价","投标文件递交", "招标范围"],
        "二、资信及业绩要求": ["资质要求", "业绩要求"],
        "三、主要人员资格要求、财务要求、信誉要求": ["人员、财务与信誉要求", "其他要求"],
        "四、踏勘现场": ["踏勘现场"],
        "五、最高投标限价": ["最高投标限价"],
        "六、评标方法": ["摘要内容-评标办法"],
        "七、合同价格与支付": ["合同价格与支付"],
        "八、罚则": ["罚则"],
        "九、其他合同条件": [],
        "十、投标报价": []
    """
    doc = Document()

    style_title = doc.styles.add_style('style_title', WD_STYLE_TYPE.PARAGRAPH)
    style_title.font.name = 'Arial'
    style_title.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    style_title.font.size = Pt(15)
    style_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_title.paragraph_format.line_spacing = 1.5

    style_level1 = doc.styles.add_style('style_level1', WD_STYLE_TYPE.PARAGRAPH)
    style_level1.base_style = doc.styles['Heading 1']  # 基于现有的标题1样式
    style_level1.font.name = 'Arial'  # 拉丁文本字体
    style_level1.element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')  # 亚洲文本字体
    style_level1.font.bold = True
    style_level1.font.size = Pt(14)
    style_level1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style_level1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_level1.paragraph_format.outline_level = 0

    style_level2 = doc.styles.add_style('style_level2', WD_STYLE_TYPE.PARAGRAPH)
    style_level2.font.name = 'Arial'
    style_level2.font.size = Pt(14)
    style_level2.font.bold = True
    style_level2.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
    style_level2.paragraph_format.line_spacing = 1.5  # 1.5行间距
    style_level2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_level2.paragraph_format.outline_level = 1  # 大纲级别为2时，outline_level设为1
    
    style = doc.styles.add_style('style', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.first_line_indent = Inches(0.5)

    style_reference = doc.styles.add_style('style_reference', WD_STYLE_TYPE.PARAGRAPH)
    style_reference.font.name = 'Arial'
    style_reference.element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong')
    style_reference.font.underline = True
    style_reference.font.size = Pt(14)
    style_reference.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style_reference.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_reference.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_paragraph(f"{title}\n{doc_type}", style='style_title')
    for section, subsections in outline.items():
        doc.add_paragraph(section, style='style_level1')
        if subsections:
            if isinstance(subsections, list):
                for subsection in subsections:
                        doc.add_paragraph(subsection, style='style_level2')
                        if dp:=element_json.get(subsection, False):
                            if dp.get("conclusion", False):
                                 doc.add_paragraph(dp["conclusion"], style='style')
                            if dp.get("tender_source", False):
                                 doc.add_paragraph(dp["tender_source"], style='style')
                            for line in re.split(r'\n|\\n', element_json[subsection]["content"]):
                                doc.add_paragraph(line, style='style')
                            if dp.get('reference', False):
                                doc.add_paragraph(element_json[subsection]["reference"], style='style_reference')
            if isinstance(subsections, str):
                for line in re.split(r'\n|\\n', subsections):
                    doc.add_paragraph(line, style='style')
                    

    doc.save(docx_path)